from __future__ import annotations

import base64
import html
import json
import re
import sqlite3
from datetime import datetime, timedelta, timezone
from io import BytesIO
from typing import cast
from uuid import uuid4

from docx import Document as WordDocument
from fastapi import APIRouter, BackgroundTasks, Depends, Response, status
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from ..config import Settings, get_settings
from ..database import create_connection
from ..deps import get_current_user, get_db
from ..document_access import ensure_document_access
from ..errors import api_error
from ..schemas import CreateExportRequest, ExportJobResponse, ExportJobStatusResponse, ReadyExportResponse

router = APIRouter(tags=["exports"])

EXPORT_CONTENT_TYPES = {
    "txt": "text/plain; charset=utf-8",
    "json": "application/json; charset=utf-8",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def html_to_plain_text(value: str) -> str:
    normalized = re.sub(r"<br\s*/?>", "\n", value, flags=re.IGNORECASE)
    normalized = re.sub(r"</(p|div|h1|h2|h3|li|pre|blockquote)>", "\n", normalized, flags=re.IGNORECASE)
    normalized = re.sub(r"<(li)[^>]*>", "• ", normalized, flags=re.IGNORECASE)
    normalized = re.sub(r"<[^>]+>", "", normalized)
    normalized = html.unescape(normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def render_pdf_bytes(document: dict) -> bytes:
    buffer = BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=A4, title=document["title"], author=document.get("owner_user_id", "Collaborative Editor AI"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ExportTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, leading=22, spaceAfter=14)
    body_style = ParagraphStyle("ExportBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=11, leading=15, spaceAfter=8)
    meta_style = ParagraphStyle("ExportMeta", parent=styles["BodyText"], fontName="Helvetica-Oblique", fontSize=9, textColor="#5f6368", spaceAfter=12)

    story = [
        Paragraph(html.escape(document["title"]), title_style),
        Paragraph(f"Updated {html.escape(document['updated_at'])}", meta_style),
    ]

    plain_text = html_to_plain_text(document["content"])
    paragraphs = [segment.strip() for segment in re.split(r"\n\s*\n", plain_text) if segment.strip()]
    if not paragraphs:
        paragraphs = ["(empty document)"]

    for paragraph in paragraphs:
        safe_text = html.escape(paragraph).replace("\n", "<br/>")
        story.append(Paragraph(safe_text, body_style))
        story.append(Spacer(1, 6))

    pdf.build(story)
    return buffer.getvalue()


def render_docx_bytes(document: dict) -> bytes:
    word_document = WordDocument()
    word_document.core_properties.title = document["title"]
    word_document.add_heading(document["title"], level=1)
    word_document.add_paragraph(f"Updated {document['updated_at']}")

    plain_text = html_to_plain_text(document["content"])
    paragraphs = [segment.strip() for segment in re.split(r"\n\s*\n", plain_text) if segment.strip()]
    if not paragraphs:
        paragraphs = ["(empty document)"]

    for paragraph_text in paragraphs:
        word_document.add_paragraph(paragraph_text)

    buffer = BytesIO()
    word_document.save(buffer)
    return buffer.getvalue()


def create_export_payload(document: dict, export_format: str, job_id: str) -> dict:
    expires_at = (datetime.now(timezone.utc) + timedelta(minutes=15)).isoformat()
    if export_format == "json":
        content = json.dumps(
            {
                "documentId": document["document_id"],
                "title": document["title"],
                "content": document["content"],
                "updatedAt": document["updated_at"],
            },
            indent=2,
        )
    elif export_format == "pdf":
        content = base64.b64encode(render_pdf_bytes(document)).decode("ascii")
    elif export_format == "docx":
        content = base64.b64encode(render_docx_bytes(document)).decode("ascii")
    else:
        content = f"Export format: {export_format}\nTitle: {document['title']}\n\n{document['content']}"

    payload = {
        "downloadUrl": f"/exports/{job_id}/download",
        "expiresAt": expires_at,
        "content": content,
        "contentType": EXPORT_CONTENT_TYPES[export_format],
        "fileName": f"{document['document_id']}.{export_format}",
    }
    if export_format in {"pdf", "docx"}:
        payload["contentEncoding"] = "base64"
    return payload


def finalize_export_job(settings: Settings, job_id: str, document: dict, export_format: str) -> None:
    connection = create_connection(settings)
    try:
        timestamp = datetime.now(timezone.utc).isoformat()
        connection.execute(
            "UPDATE export_jobs SET status = 'RUNNING', updated_at = ? WHERE job_id = ?",
            (timestamp, job_id),
        )
        result = create_export_payload(document, export_format, job_id)
        connection.execute(
            "UPDATE export_jobs SET status = 'SUCCEEDED', result_json = ?, updated_at = ? WHERE job_id = ?",
            (json.dumps(result), timestamp, job_id),
        )
        connection.commit()
    finally:
        connection.close()


@router.post("/documents/{document_id}/export", response_model=ReadyExportResponse | ExportJobResponse, status_code=status.HTTP_200_OK)
def create_export(
    document_id: str,
    payload: CreateExportRequest,
    background_tasks: BackgroundTasks,
    current_user=Depends(get_current_user),
    db: sqlite3.Connection = Depends(get_db),
    settings: Settings = Depends(get_settings),
):
    document = ensure_document_access(db, document_id, current_user["user_id"], "viewer")
    if payload.format not in EXPORT_CONTENT_TYPES:
        raise api_error(400, "INVALID_INPUT", "unsupported export format")

    if payload.format in {"txt", "json"}:
        return ReadyExportResponse.model_validate(create_export_payload(document, payload.format, f"expjob_{uuid4().hex[:12]}"))

    job_id = f"expjob_{uuid4().hex[:12]}"
    timestamp = datetime.now(timezone.utc).isoformat()
    db.execute(
        "INSERT INTO export_jobs (job_id, document_id, user_id, format, status, result_json, error_code, error_message, created_at, updated_at) VALUES (?, ?, ?, ?, 'PENDING', NULL, NULL, NULL, ?, ?)",
        (job_id, document_id, current_user["user_id"], payload.format, timestamp, timestamp),
    )
    db.commit()
    background_tasks.add_task(finalize_export_job, settings, job_id, document, payload.format)
    return ExportJobResponse(job_id=job_id, status_url=f"/exports/{job_id}")


@router.get("/exports/{job_id}", response_model=ExportJobStatusResponse)
def get_export_status(job_id: str, current_user=Depends(get_current_user), db: sqlite3.Connection = Depends(get_db)) -> ExportJobStatusResponse:
    row = db.execute(
        "SELECT job_id, document_id, status, result_json, error_code, error_message FROM export_jobs WHERE job_id = ?",
        (job_id,),
    ).fetchone()
    if row is None:
        raise api_error(404, "NOT_FOUND", "export job not found")
    ensure_document_access(db, row["document_id"], current_user["user_id"], "viewer")
    result = json.loads(row["result_json"]) if row["result_json"] else {}
    return ExportJobStatusResponse(
        job_id=row["job_id"],
        status=row["status"],
        download_url=result.get("downloadUrl"),
        expires_at=result.get("expiresAt"),
        error_code=row["error_code"],
        error_message=row["error_message"],
    )


@router.get("/exports/{job_id}/download")
def download_export(job_id: str, current_user=Depends(get_current_user), db: sqlite3.Connection = Depends(get_db)) -> Response:
    row = db.execute(
        "SELECT job_id, document_id, status, result_json FROM export_jobs WHERE job_id = ?",
        (job_id,),
    ).fetchone()
    if row is None:
        raise api_error(404, "NOT_FOUND", "export job not found")
    ensure_document_access(db, row["document_id"], current_user["user_id"], "viewer")
    if row["status"] != "SUCCEEDED" or not row["result_json"]:
        raise api_error(409, "CONFLICT", "export is not ready for download")
    result = cast(dict, json.loads(row["result_json"]))
    content = result["content"]
    if result.get("contentEncoding") == "base64":
        content = base64.b64decode(content)
    return Response(
        content=content,
        media_type=result["contentType"],
        headers={"Content-Disposition": f"attachment; filename=\"{result['fileName']}\""},
    )
